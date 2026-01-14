import os
import json
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from flask import render_template, redirect, url_for, flash, session
import re
from google import genai
from authlib.integrations.flask_client import OAuth
from flask import Flask, request, jsonify, send_file, send_from_directory
from datetime import datetime
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from datetime import timedelta
from functools import wraps
import io
import zipfile
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask_cors import CORS
import fitz  # PyMuPDF
from flask import Flask, request, render_template
from twilio.twiml.messaging_response import MessagingResponse
from flask import request, jsonify, send_file
from PIL import Image, ImageDraw, ImageFont
import textwrap
import uuid
import PyPDF2  
import docx as docx_reader 
import csv

def is_combo_user(user):
    return (
        user.grade == "10-12"
        or user.board == "CBSE-ICSE"
    )


EXAM_STRUCTURES = {
    "SSC CGL": {
        "Quantitative Aptitude": 25,
        "Logical Reasoning": 25,
        "English Language": 25,
        "General Awareness": 25
    },
    "IBPS PO": {
        "Quantitative Aptitude": 35,
        "Reasoning Ability": 35,
        "English Language": 30
    },
    "SBI PO": {
        "Quantitative Aptitude": 35,
        "Reasoning Ability": 35,
        "English Language": 30
    },
    "RRB NTPC": {
        "Mathematics": 30,
        "General Intelligence & Reasoning": 30,
        "General Awareness": 40
    },
    "UPSC NDA": {
        "Mathematics": 120,
        "General Ability Test": 150
    },
    "GATE": {
        "Engineering Mathematics": 15,
        "General Aptitude": 15,
        "Core Subject": 70
    },
    "IBPS SO": {
    "Professional Knowledge": 50,
    "Reasoning Ability": 50
}

}


BASE_DIR = os.path.dirname(os.path.abspath(__file__))
db_path = os.path.join(BASE_DIR, "users.db")

GLOBAL_HEADER = "MathGen • Practice Worksheets"
GLOBAL_FOOTER = "© MathGen | For educational use only"


TEMP_DIR = os.path.join(BASE_DIR, "temp_files")
os.makedirs(TEMP_DIR, exist_ok=True)


def is_valid_indian_mobile(phone):
    # Must be exactly 10 digits and start with 6–9
    if not re.fullmatch(r"[6-9]\d{9}", phone):
        return False

    # Reject repeated digits (0000000000, 1111111111, etc.)
    if len(set(phone)) == 1:
        return False

    # Reject common fake numbers
    if phone in {"1234567890", "0123456789", "9876543210"}:
        return False

    return True

TOPICS_CSV = os.path.join(BASE_DIR, "topics.csv")


app = Flask(__name__,
            static_folder='static',
            template_folder='templates')

CORS(app, supports_credentials=True)

app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE='Lax',   # ✅ SAFE default
    SESSION_COOKIE_SECURE=False,     # ✅ local dev
    REMEMBER_COOKIE_DURATION=0,
    PERMANENT_SESSION_LIFETIME=timedelta(hours=3)
)
# 🔐 Production-only cookie fix (Render)
if os.environ.get("RENDER") == "true":

    app.config.update(
        SESSION_COOKIE_SECURE=True   # ✅ HTTPS on Render
    )

app.config['SECRET_KEY'] = os.environ.get('FLASK_SECRET_KEY', 'default-super-secret-key-change-me-immediately')
app.config['GOOGLE_CLIENT_ID'] = os.environ.get('GOOGLE_CLIENT_ID')
app.config['GOOGLE_CLIENT_SECRET'] = os.environ.get('GOOGLE_CLIENT_SECRET')
app.config['GENAI_API_KEY'] = os.environ.get('GENAI_API_KEY')
app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{db_path}'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['GOOGLE_CONF_URL'] = 'https://accounts.google.com/.well-known/openid-configuration'

db = SQLAlchemy(app)

@app.context_processor
def inject_now():
    return {'now': datetime.now}

class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    country_code = db.Column(db.String(5), nullable=False, default="+91")  # ✅ ADD
    phone_number = db.Column(db.String(30), unique=True, nullable=True)
    grade = db.Column(db.String(20))
    password_hash = db.Column(db.String(256), nullable=True)
    age = db.Column(db.Integer, nullable=True)
    city = db.Column(db.String(100), nullable=True)
    postal_code = db.Column(db.String(20), nullable=True)
    timezone = db.Column(db.String(50), nullable=True)
    whatsapp_consent = db.Column(db.Boolean, default=False)
    newsletter_consent = db.Column(db.Boolean, default=False)
    board = db.Column(db.String(50), nullable=True)
    profile_completed = db.Column(db.Boolean, default=False)
    dob = db.Column(db.Date, nullable=True)


    def set_password(self, password):
        self.password_hash = generate_password_hash(password, method='pbkdf2:sha256', salt_length=16)

    def check_password(self, password):
        return self.password_hash and check_password_hash(self.password_hash, password)

    def __repr__(self):
        return f'<User {self.email}>'

class MockTest(db.Model):
    __tablename__ = "mock_test"

    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=True) # <--- ADD THIS
    title = db.Column(db.String(255))
    description = db.Column(db.Text)
    category = db.Column(db.String(50))
    duration_minutes = db.Column(db.Integer)


class MockQuestion(db.Model):
    __tablename__ = "mock_question"

    id = db.Column(db.Integer, primary_key=True)
    test_id = db.Column(db.Integer, db.ForeignKey("mock_test.id"), nullable=False)
    qno = db.Column(db.Integer)
    question_text = db.Column(db.Text)
    options_json = db.Column(db.Text)
    correct_option_index = db.Column(db.Integer)
    explanation = db.Column(db.Text)

    # ✅ ADD THESE
    def options(self):
        return json.loads(self.options_json)

    @property
    def correct_option(self):
        return self.options()[self.correct_option_index]

class MockAnswer(db.Model):
    __tablename__ = "mock_answer"

    id = db.Column(db.Integer, primary_key=True)
    attempt_id = db.Column(db.Integer, db.ForeignKey("mock_attempt.id"), nullable=False)
    question_id = db.Column(db.Integer, db.ForeignKey("mock_question.id"), nullable=False)
    selected_option = db.Column(db.Integer)

    attempt = db.relationship("MockAttempt", backref="answers")



class MockAttempt(db.Model):
    __tablename__ = "mock_attempt"

    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)
    test_id = db.Column(db.Integer, db.ForeignKey("mock_test.id"), nullable=False)

    score = db.Column(db.Integer, default=0)
    total = db.Column(db.Integer, default=0)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    test = db.relationship("MockTest")  # ✅ ADD THIS LINE



def load_topics():
    topics = []
    with open('topics.csv', newline='', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        for row in reader:
            topics.append(row)
    return topics



# ---- MOCK TEST DATA MODEL ----


@app.route('/competitive-exams')
@login_required
def competitive_exams():
    tests = MockTest.query.filter_by(category='competitive').all()
    return render_template('competitive.html', tests=tests)


@app.route('/job-exams')
@login_required
def job_exams():
    daily_tests = MockTest.query.filter_by(category='job_daily').all()
    weekly_tests = MockTest.query.filter_by(category='job_weekly').all()
    job_tests = MockTest.query.filter_by(category='job').all()

    return render_template(
        'job_exams.html',
        daily_tests=daily_tests,
        weekly_tests=weekly_tests,
        job_tests=job_tests,
        student_name=current_user.name,
        student_grade=current_user.grade,
        student_board=current_user.board
    )


def generate_gemini_questions(topic, count):

    # ✅ correct client creation (NO configure)
    client = genai.Client(
    api_key=os.environ.get("GENAI_API_KEY"))

    prompt = f"""
You are an exam question generator.

Generate {count} multiple-choice questions on:
Topic: {topic}

Rules:
- Each question must have 4 options
- Clearly specify the correct option index (0–3)
- Provide a short explanation

IMPORTANT:
- Respond with ONLY a valid JSON array
- Do NOT include explanations
- Do NOT include markdown
- Do NOT include extra text
CRITICAL FORMAT RULES (MANDATORY):
- Use ONLY plain text math
- NO LaTeX
- NO $, \\, \\alpha, \\beta, \\sqrt
- Use unicode symbols ONLY:
  π, α, β, √, ≤, ≥, ×, ÷
- Use normal brackets only: (), [], {{}},
- Write math exactly how students write in notebooks

[
  {{
    "question": "...",
    "options": ["A", "B", "C", "D"],
    "correct": 1,
    "explanation": "..."
  }}
]
"""

    response = client.models.generate_content(
        model="models/gemini-flash-latest",
        contents=prompt
    )

    try:
        raw = extract_json_from_ai(response.text)
        if not raw:
            return []

        for q in raw:
            q["question"] = clean_ai_text(q["question"])
            q["options"] = [clean_ai_text(opt) for opt in q["options"]]
            q["explanation"] = clean_ai_text(q["explanation"])

        return raw

    except Exception as e:
        print("❌ Gemini JSON parse error:", e)
        print("RAW RESPONSE ↓↓↓")
        print(response.text)
        return []


def start_job_ai_test(job_type, exam_type, authority, past_paper, test_type, count):
    """
    AI engine for job exams:
    A) Solve past paper
    B) Generate mock questions
    C) Evaluate later
    """

    client = genai.Client(api_key=os.environ.get("GENAI_API_KEY"))

    # ------------------------------
    # A) SOLUTION GENERATION
    # ------------------------------
    solutions = []
    if past_paper:
        solution_prompt = f"""
You are an expert examiner for {job_type.upper()} exams.

Exam Type: {exam_type}
Authority: {authority}

Solve the following questions.
Provide:
- Question
- Correct answer
- Short explanation

Return STRICT JSON ONLY:

[
  {{
    "question": "...",
    "answer": "...",
    "explanation": "..."
  }}
]

Questions:
{past_paper}
"""
        sol_resp = client.models.generate_content(
            model="models/gemini-flash-latest",
            contents=solution_prompt
        )

        solutions = extract_json_from_ai(sol_resp.text) or []

    # ------------------------------
    # B) MOCK TEST GENERATION
    # ------------------------------
    job_type_upper = job_type.upper() if job_type else "JOB"

    mock_prompt = f"""
You are generating a mock test for {job_type_upper} exams.

Exam Type: {exam_type}
Authority: {authority}

Generate {count} MCQs.
Rules:
- 4 options
- Give correct option index (0-3)
- Difficulty similar to real exam
- STRICT JSON ONLY

[
  {{
    "question": "...",
    "options": ["A","B","C","D"],
    "correct": 0,
    "explanation": "..."
  }}
]
"""

    mock_resp = client.models.generate_content(
        model="models/gemini-flash-latest",
        contents=mock_prompt
    )

    questions = extract_json_from_ai(mock_resp.text)

    if not questions:
        flash("AI could not generate job test questions.", "danger")
        return redirect(url_for("job_exams"))

    # ------------------------------
    # SAVE TEST
    # ------------------------------
    title = f"{job_type_upper} {exam_type.upper()} Mock Test"

    test = MockTest(
        title=title,
        category=test_type,
        duration_minutes=20
    )
    db.session.add(test)
    db.session.flush()

    for idx, q in enumerate(questions, start=1):
        mq = MockQuestion(
            test_id=test.id,
            qno=idx,
            question_text=clean_ai_text(q["question"]),
            options_json=json.dumps(q["options"]),
            correct_option_index=int(q["correct"]),
            explanation=clean_ai_text(q["explanation"])
        )
        db.session.add(mq)

    attempt = MockAttempt(
        user_id=current_user.id,
        test_id=test.id,
        total=len(questions)
    )
    db.session.add(attempt)
    db.session.commit()

    return redirect(url_for("take_test", attempt_id=attempt.id))

@app.route("/jobseekers", methods=["GET"])
@login_required   # optional but recommended
def jobseekers():
    return render_template("job_seekers.html")


@app.route("/start-test", methods=["POST"])
@login_required
def start_test():
    uploaded_file = request.files.get("worksheet_file")
    past_paper_text = request.form.get("past_paper", "").strip()
    job_type = request.form.get("job_type")
    exam_type = request.form.get("exam_type")
    exam_authority = request.form.get("exam_authority")
    test_type = request.form.get("test_type", "daily")
    topic = request.form.get("topic")                     # MAJOR topic
    minor_topic = request.form.get("minor_topic")         # ✅ ADD (OPTIONAL)
    count = int(request.form.get("question_count", 5))
    
    # --------------------------------------
    # READ PAST PAPER FROM FILE (IF UPLOADED)
    # --------------------------------------
    if uploaded_file and uploaded_file.filename:
        filename = uploaded_file.filename.lower()

        if filename.endswith(".pdf"):
            extracted = get_text_from_pdf(uploaded_file)

        elif filename.endswith(".docx"):
            extracted = get_text_from_docx(uploaded_file)

        elif filename.endswith(".txt"):
            extracted = uploaded_file.read().decode("utf-8", errors="ignore")

        else:
            extracted = None

        if extracted:
            past_paper_text = extracted

    
    # ======================================================
    # JOB-BASED AI TEST FLOW (Manager Requirement)
    # ======================================================
    if job_type:
        if not exam_type or not exam_authority:
            flash("Please select exam type and authority.", "warning")
            return redirect(url_for("job_exams"))

        return start_job_ai_test(
            job_type=job_type,
            exam_type=exam_type,
            authority=exam_authority,
            past_paper=past_paper_text,
            test_type=test_type,
            count=count
        )


    # 🚨 HARD VALIDATION (MUST)
    if not topic:
        flash("Please select a topic/job type before starting the test.", "warning")
        return redirect(url_for("job_exams"))

    # ✅ FINAL TOPIC SELECTION (SAFE FALLBACK)
    final_topic = minor_topic if minor_topic else topic   # ✅ ADD

    try:
        questions = generate_gemini_questions(final_topic, count)
        if len(questions) < count:
            flash(f"AI generated only {len(questions)} questions. Try again.", "warning")
            return redirect(url_for("job_exams"))

    except Exception as e:
        print("❌ Gemini error:", e)
        flash("AI service is temporarily unavailable. Please try again later.", "danger")
        return redirect(url_for("job_exams"))

    # 🚨 EMPTY RESULT CHECK
    if not questions or len(questions) == 0:
        flash("Could not generate questions. Try again after some time.", "danger")
        return redirect(url_for("job_exams"))

    # ✅ Create Test
    # Inside start_test:
    test = MockTest(
        title=f"{final_topic} Test",
        user_id=current_user.id,  # <--- SECURE THIS TEST TO THE USER
        category=test_type,
        duration_minutes=10
    )
    db.session.add(test)
    db.session.commit() # Commit here to get the ID safely

    # ✅ Save questions safely
    for idx, q in enumerate(questions, start=1):
        if not all(k in q for k in ("question", "options", "correct", "explanation")):
            continue

        mq = MockQuestion(
            test_id=test.id,
            qno=idx,
            question_text=q["question"],
            options_json=json.dumps(q["options"]),
            correct_option_index=int(q["correct"]),
            explanation=q["explanation"]
        )
        db.session.add(mq)

    # ✅ Create attempt
    attempt = MockAttempt(
        user_id=current_user.id,
        test_id=test.id,
        total=len(questions)
    )
    db.session.add(attempt)
    db.session.commit()

    # ✅ THIS REDIRECT WILL NOW ALWAYS WORK
    return redirect(url_for("take_test", attempt_id=attempt.id))




@app.route('/submit-test/<int:attempt_id>', methods=['POST'])
@login_required
def submit_test(attempt_id):
    attempt = MockAttempt.query.get_or_404(attempt_id)
    questions = MockQuestion.query.filter_by(test_id=attempt.test_id).all()

    score = 0

    for q in questions:
        selected = request.form.get(f"q_{q.id}")
        if selected is None:
            continue

        selected = int(selected)

        ans = MockAnswer(
            attempt_id=attempt.id,
            question_id=q.id,
            selected_option=selected
        )
        db.session.add(ans)

        if selected == q.correct_option_index:
            score += 1

    attempt.score = score
    attempt.total = len(questions)
    db.session.commit()

    return redirect(url_for("my_scores"))

@app.route("/get-topics")
@login_required
def get_topics():
    board = request.args.get("board", "").strip().lower()
    grade = request.args.get("grade", "").strip()
    subject = "mathematics"

    # ---------- HANDLE COMBO GRADE ----------
    if grade == "10-12":
        grades = {"10", "11", "12"}
    else:
        grades = {grade}

    # ---------- HANDLE COMBO BOARD ----------
    if board == "cbse-icse":
        boards = {"cbse", "icse"}
    else:
        boards = {board}

    major_topics = set()

    with open(TOPICS_CSV, newline="", encoding="utf-8", errors="ignore") as f:
        reader = csv.DictReader(f)

        for row in reader:
            row_board = row.get("board", "").strip().lower()
            row_grade = row.get("grade", "").strip()
            row_subject = row.get("subject", "").strip().lower()
            row_major = row.get("major_topic", "").strip()

            if (
                row_board in boards
                and row_grade in grades
                and row_subject == subject
                and row_major
            ):
                major_topics.add(row_major)

    return jsonify(sorted(list(major_topics)))



@app.route("/get-minor-topics")
@login_required
def get_minor_topics():
    board = request.args.get("board", "").strip().lower()
    grade = str(request.args.get("grade", "")).strip()
    subject = request.args.get("subject", "Mathematics").strip().lower()
    major_topic = request.args.get("major_topic", "").strip()

    minor_topics = set()

    if not major_topic:
        return jsonify([])

    with open(TOPICS_CSV, newline="", encoding="utf-8", errors="ignore") as f:
        reader = csv.DictReader(f)

        for row in reader:
            row_board = row.get("board", "").strip().lower()
            row_grade = row.get("grade", "").strip()
            row_subject = row.get("subject", "").strip().lower()
            row_major = row.get("major_topic", "").strip()
            row_minor = row.get("minor_topic", "").strip()

            if grade == "10-12":
                grades = {"10", "11", "12"}
            else:
                grades = {grade}

            if board == "cbse-icse":
                boards = {"cbse", "icse"}
            else:
                boards = {board}

            if (
                row_board in boards
                and row_grade in grades
                and row_subject == subject
                and row_major == major_topic
                and row_minor
            ):
                minor_topics.add(row_minor)

    return jsonify(sorted(list(minor_topics)))

@app.route('/my-scores')
@login_required
def my_scores():
    attempts = (
        MockAttempt.query
        .filter_by(user_id=current_user.id)
        .order_by(MockAttempt.created_at.desc())
        .all()
    )

    rows = []
    total_score = 0
    total_questions = 0

    for idx, attempt in enumerate(attempts, start=1):
        # ✅ safer than .query.get()
        test = MockTest.query.filter_by(id=attempt.test_id).first()

        score = attempt.score or 0
        total = attempt.total or 0

        total_score += score
        total_questions += total

        rows.append({
            "sr_no": idx,
            "attempt_id": attempt.id,
            "test_title": test.title if test else "Mock Test",
            "score": score,
            "total": total,
            "date": attempt.created_at.strftime("%d %b %Y, %H:%M")
        })

    # ✅ default-safe values
    accuracy = 0
    avg_score = 0
    max_total = 0

    if total_questions > 0:
        accuracy = round((total_score / total_questions) * 100, 2)

    if rows:
        avg_score = round(total_score / len(rows), 2)
        max_total = max(row["total"] for row in rows)

    return render_template(
        "my_scores.html",
        rows=rows,
        accuracy=accuracy,
        avg_score=avg_score,
        max_total=max_total
    )


# -------------------------

# --- LOGIN MANAGER SETUP ---
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = None
login_manager.refresh_view = None




@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# --- OAuth Setup ---
oauth = OAuth(app)
oauth.register(
    name='google',
    client_id=app.config['GOOGLE_CLIENT_ID'],
    client_secret=app.config['GOOGLE_CLIENT_SECRET'],
    server_metadata_url=app.config['GOOGLE_CONF_URL'],
    client_kwargs={
        'scope': 'openid email profile'
    }
)


def create_docx(content, title, sub_title_info, filename):
    file_path = os.path.join(TEMP_DIR, filename)

    doc = Document()
    section = doc.sections[0]

    # ---------- HEADER ----------
    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = GLOBAL_HEADER
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---------- TITLE ----------
    title_p = doc.add_heading(title, level=1)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---------- SUB HEADER ----------
    sub = (
        f"Date: {sub_title_info['date']}   |   "
        f"Marks: {sub_title_info['marks']}   |   "
        f"Topic: {sub_title_info['sub-title']}"
    )
    sub_p = doc.add_paragraph(sub)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("") # Spacer

    # ---------- CONTENT (The Fix) ----------
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("") # Vertical gap
            continue
        
        p = doc.add_paragraph(line)
        # This prevents the "cramped" look and ensures vertical stacking
        p.paragraph_format.space_after = Pt(12) 
        p.paragraph_format.line_spacing = 1.5

    # ---------- FOOTER ----------
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = GLOBAL_FOOTER
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(file_path)
    return file_path

def create_txt(content, title, sub_title_info, filename):
    file_path = os.path.join(TEMP_DIR, filename)
    
    with open(file_path, "w", encoding="utf-8") as f:
        # ---------- HEADER ----------
        f.write(f"{GLOBAL_HEADER}\n")
        f.write(f"{title}\n")
        f.write("=" * 60 + "\n\n")

        f.write(f"Date  : {sub_title_info.get('date', '')}\n")
        f.write(f"Marks : {sub_title_info.get('marks', '')}\n")
        f.write(f"Topic : {sub_title_info.get('sub-title', '')}\n\n")

        f.write("-" * 60 + "\n")
        f.write("QUESTIONS / SOLUTIONS\n")
        f.write("-" * 60 + "\n\n")

        # ---------- CONTENT (The Fix) ----------
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                f.write("\n")
                continue
            
            # Manually wrap the text at 80 characters for TXT files
            wrapped_lines = textwrap.wrap(line, width=80)
            for wl in wrapped_lines:
                f.write(wl + "\n")
            f.write("\n") # Add a small gap between different items

        # ---------- FOOTER ----------
        f.write("\n" + "=" * 60 + "\n")
        f.write(GLOBAL_FOOTER + "\n")

    return file_path



# --- Helper functions to read uploaded files ---
def get_text_from_pdf(file_storage):
    """Extracts text from a text-based PDF file (FileStorage object)."""
    try:
        pdf_reader = PyPDF2.PdfReader(file_storage.stream)
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text if text else None
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return None

def get_text_from_docx(file_storage):
    """Extracts text from a DOCX file (FileStorage object)."""
    try:
        doc = docx_reader.Document(file_storage.stream)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text if text else None
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return None
    
import pytesseract

def get_text_from_image(file_storage):
    try:
        file_storage.stream.seek(0)
        image = Image.open(file_storage.stream)

        # 🔥 FORCE RGB (CRITICAL FOR PNG/GIF)
        if image.mode != "RGB":
            image = image.convert("RGB")

        text = pytesseract.image_to_string(image)
        return text.strip()

    except Exception as e:
        print("OCR ERROR:", e)
        return None


    
def create_image(content, title, sub_title_info, filename, fmt="png"):
    file_path = os.path.join(TEMP_DIR, filename)

    width, height = 1240, 1754
    margin_x, margin_y = 60, 60
    line_height = 35 # Increased slightly for better readability

    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)

    FONT_PATH = os.path.join(BASE_DIR, "fonts", "DejaVuSans.ttf")
    font_title = ImageFont.truetype(FONT_PATH, 36)
    font_body = ImageFont.truetype(FONT_PATH, 22)
    
    # --- GLOBAL HEADER ---
    draw.text((width // 2 - 220, 20), GLOBAL_HEADER, font=font_body, fill="black")

    y = margin_y
    draw.text((margin_x, y), title, font=font_title, fill="black")
    y += 60

    # --- SUBTITLE ---
    subtitle = f"Date: {sub_title_info['date']} | Marks: {sub_title_info['marks']} | Topic: {sub_title_info['sub-title']}"
    draw.text((margin_x, y), subtitle, font=font_body, fill="black")
    y += 50 # Extra space after subtitle

    # --- CONTENT LOOP (STRICT VERTICAL ALIGNMENT) ---
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            y += 20 # Add vertical gap for empty lines
            continue

        # Wrap text to ensure it doesn't go off the right side
        wrapped = textwrap.wrap(line, width=70) # Wrap at ~70 characters
        for w in wrapped:
            if y > height - 100: # Check for page bottom
                break
            draw.text((margin_x, y), w, font=font_body, fill="black")
            y += line_height # FORCE vertical increment for every wrapped line
        
        y += 10 # Add a small vertical buffer between different questions/answers

    # --- GLOBAL FOOTER ---
    draw.text((width // 2 - 260, height - 60), GLOBAL_FOOTER, font=font_body, fill="black")

    if fmt == "gif":
        img.convert("P", palette=Image.ADAPTIVE).save(file_path, format="GIF")
    else:
        img.save(file_path)

    return file_path



# ----------------------------------------------------

# --- Static File and Main Page Routes ---


@app.route('/')
def landing_page():
    show_profile_popup = (
    current_user.is_authenticated
    and not current_user.profile_completed
)


    return render_template(
        'landing.html',
        show_profile_popup=show_profile_popup
    )



@app.route('/generator')
@login_required
def serve_index():
    # 🔒 Absolute rule for combo users
    if is_combo_user(current_user):
        return redirect(url_for("exam_combo_page"))

    return render_template(
        "index.html",
        user_grade=current_user.grade,
        user_board=current_user.board
    )






@app.route('/static/<path:filename>')
def serve_static(filename):
     return send_from_directory(app.static_folder, filename)

# --- NEW: Static Page Routes ---
@app.route('/about')
@login_required # Protect this page
def about():
    return render_template('about.html')

@app.route('/features')
@login_required # Protect this page
def features():
    return render_template('features.html')
# --- END NEW ROUTES ---


# --- AUTHENTICATION ROUTES ---
from datetime import datetime

@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('serve_index'))

    errors = {}
    form = {}

    if request.method == 'POST':
        form = request.form.to_dict()

        # ---------- READ INPUT ----------
        name = form.get('name', '').strip()
        email = form.get('email', '').strip()
        phone_number_main = form.get('phone_number_main', '').strip()
        grade = form.get('grade', '').strip()
        board = form.get('board', '').strip()
        age = form.get('age', '')
        city = form.get('city', '')
        postal_code = form.get('postal_code', '')
        dob_str = form.get('dob', '').strip()
        password = form.get('password', '')
        confirm_password = form.get('confirm_password', '')
        country_code = form.get('country_code', '+91').strip()

        # ---------- REQUIRED ----------
        if not name:
            errors['name'] = 'Full name is required.'

        if not email:
            errors['email'] = 'Email is required.'

        if not phone_number_main:
            errors['phone'] = 'Phone number is required.'

        if not grade:
            errors['grade'] = 'Grade is required.'
        
        if not age:
            errors['age'] = 'Age is required.'

        if not city:
            errors['city'] = 'City is required.'

        if not postal_code:
            errors['postal_code'] = 'Postal code is required.'

        if not board:
            errors['board'] = 'Board is required.'

        if not password:
            errors['password'] = 'Password is required.'

        if not confirm_password:
            errors['confirm_password'] = 'Please confirm your password.'

        # ---------- GRADE ----------
        # ---------- GRADE ----------
        if grade:
            if grade == "10-12":
                pass  # ✅ combo allowed
            elif grade.isdigit() and 1 <= int(grade) <= 12:
                pass
            else:
                errors['grade'] = 'Invalid grade selection.'

                
        if age:
            if not age.isdigit() or int(age) < 1 or int(age) > 100:
                errors['age'] = 'Enter a valid age.'


        # ---------- DOB (CRITICAL FIX) ----------
        dob = None
        if dob_str:
            try:
                dob = datetime.strptime(dob_str, "%Y-%m-%d").date()
            except ValueError:
                errors['dob'] = 'Invalid date format.'

        # ---------- PHONE ----------
        # ---------- PHONE (STRICT VALIDATION) ----------
        phone_digits = re.sub(r"\D", "", phone_number_main)

        if not phone_digits:
            errors['phone'] = 'Mobile number is required.'

        elif not is_valid_indian_mobile(phone_digits):
            errors['phone'] = 'Enter a valid 10-digit Indian mobile number.'

        full_phone = f"{country_code}{phone_digits}"


        # ---------- EMAIL ----------
        email_regex = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        if email and not re.match(email_regex, email):
            errors['email'] = 'Invalid email format.'

        # ---------- PASSWORD ----------
        if password:
            if len(password) < 8:
                errors['password'] = 'Must be at least 8 characters.'
            elif not re.search(r'[A-Z]', password):
                errors['password'] = 'Must contain an uppercase letter.'
            elif not re.search(r'[a-z]', password):
                errors['password'] = 'Must contain a lowercase letter.'
            elif not re.search(r'\d', password):
                errors['password'] = 'Must contain a number.'
            elif not re.search(r'[!@#$%^&*()_+=\-\[\]{};:\'",.<>/?~`]', password):
                errors['password'] = 'Must contain a special character.'

        if password and confirm_password and password != confirm_password:
            errors['confirm_password'] = 'Passwords do not match.'

        # ---------- UNIQUENESS ----------
        if email and User.query.filter_by(email=email).first():
            errors['email'] = 'Email already registered.'

        if phone_digits and User.query.filter_by(phone_number=full_phone).first():
            errors['phone'] = 'Phone number already registered.'

        # ---------- FIELD ERRORS ----------
        if errors:
            return render_template('register.html', errors=errors, form=form)

        # ---------- CREATE USER ----------
        whatsapp_consent = 'whatsapp_consent' in request.form
        newsletter_consent = 'newsletter_consent' in request.form
        timezone = form.get('timezone')
        try:
            new_user = User(
                name=name,
                email=email,
                phone_number=full_phone,
                country_code=country_code,
                grade=grade,  # ✅ KEEP AS STRING
                board=board,
                age=int(age),                 # ✅ ADD  
                dob=dob,                # ✅ ADD
                city=city,                    # ✅ ADD
                postal_code=postal_code,
                timezone=timezone, 
                whatsapp_consent=whatsapp_consent,
                newsletter_consent=newsletter_consent,
                profile_completed=True
            )

            new_user.set_password(password)
            db.session.add(new_user)
            db.session.commit()

            flash('Account created successfully. Please login.', 'success')
            return redirect(url_for('landing_page', show_login=1))




        except Exception as e:
            db.session.rollback()
            print("❌ REGISTRATION ERROR:", e)
            errors['general'] = 'Unexpected server error. Please try again.'
            return render_template('register.html', errors=errors, form=form)

    return render_template('register.html', errors={}, form={})


from flask import jsonify

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return jsonify({
            "success": True,
            "redirect": url_for("serve_index")
        })


    login_method = request.form.get("login_method")
    login_identifier = request.form.get("login_identifier", "").strip()
    password = request.form.get("password")

    if login_method != "google" and (not login_identifier or not password):
        return jsonify({
            "success": False,
            "error": "Please enter both email/phone and password."
        })

    user = None

    if login_method == "email":
        user = User.query.filter_by(email=login_identifier.lower()).first()

    elif login_method == "phone":
        phone = re.sub(r"\D", "", login_identifier)
        user = User.query.filter_by(phone_number="+91" + phone).first()

    if not user or not user.check_password(password):
        return jsonify({
            "success": False,
            "error": "Invalid email/phone or password."
        })

    login_user(user, remember=False)

    redirect_url = (
        url_for("exam_combo_page")
        if is_combo_user(user)
        else url_for("serve_index")
    )

    return jsonify({
        "success": True,
        "redirect": redirect_url
    })



@login_manager.unauthorized_handler
def unauthorized():
    if request.headers.get("X-Requested-With") == "XMLHttpRequest":
        return jsonify({"unauthorized": True}), 401
    return redirect(url_for("landing_page"))



@app.route("/login-popup")
def login_popup():
    if current_user.is_authenticated:
        return ""

    return render_template("login.html")

@app.before_request
def fix_google_users():
    if current_user.is_authenticated:
        if current_user.password_hash is None and not current_user.profile_completed:
            # Google user who hasn’t completed profile
            pass



@app.route('/profile-fragment')
@login_required
def profile_fragment():
    return render_template('partials/profile_fragment.html')


@app.route('/logout')
@login_required
def logout():
    logout_user()
    session.clear()  # Clear everything in the session
    flash("You have been logged out.", "info")
    return redirect(url_for('landing_page'))




# --- GOOGLE OAUTH LOGIN ROUTES ---
@app.route('/login/google')
def login_google():
    session.pop("show_profile_popup", None)
    session.pop("google_new_user", None)

    if not app.config.get('GOOGLE_CLIENT_ID') or not app.config.get('GOOGLE_CLIENT_SECRET'):
        return redirect(url_for('landing_page', show_login=1))

    return oauth.google.authorize_redirect(
        url_for('google_callback', _external=True)
    )


@app.route('/auth/google/callback')
def google_callback():
    try:
        token = oauth.google.authorize_access_token()
        resp = oauth.google.get('https://openidconnect.googleapis.com/v1/userinfo')
        resp.raise_for_status()
        user_info = resp.json()

        email = user_info.get('email')
        name = user_info.get('name')

        user = User.query.filter_by(email=email).first()

        # ✅ FIRST-TIME GOOGLE USER
        if not user:
            user = User(
                email=email,
                name=name,
                profile_completed=False
            )
            db.session.add(user)
            db.session.commit()

        login_user(user)

        # ✅ PROFILE NOT COMPLETED → LANDING PAGE (POPUP WILL SHOW)
        if not user.profile_completed:
            return redirect(url_for("landing_page"))


        # ✅ PROFILE COMPLETED → NORMAL FLOW
        if is_combo_user(user):
            return redirect(url_for("exam_combo_page"))
        
        # EXISTING GOOGLE USER WITH COMPLETE PROFILE
        if user and user.profile_completed:
            login_user(user)
            return redirect(url_for("serve_index"))


        return redirect(url_for("serve_index"))

    except Exception as e:
        print("GOOGLE AUTH ERROR:", e)
        return redirect(url_for('landing_page', show_login=1))





from datetime import datetime



# --- PROFILE AND SETTINGS ROUTES ---

@app.route('/profile', methods=['GET','POST'])
@login_required
def profile():
    from datetime import datetime
    import re

    is_ajax = request.headers.get("X-Requested-With") == "XMLHttpRequest"

    if request.method == "POST":

        name = request.form.get("name", "").strip()
        grade = request.form.get("grade", "").strip()
        board = request.form.get("board", "").strip()
        age = request.form.get("age", "").strip()
        city = request.form.get("city", "").strip()
        postal_code = request.form.get("postal_code", "").strip()
        timezone = request.form.get("timezone", "").strip()
        dob_str = request.form.get("dob", "").strip()
        phone_number = request.form.get("phone_number_main", "").strip()

        if not phone_number:
            return jsonify(success=False, error="Mobile number is required.")



        whatsapp_consent = bool(request.form.get("whatsapp_consent"))
        newsletter_consent = bool(request.form.get("newsletter_consent"))

        # ---------- REQUIRED ----------
        if not name or not grade or not age:
            if is_ajax:
                return jsonify(success=False, error="Please fill all required fields.")
            flash("Please fill all required fields marked with *", "danger")
            return redirect(url_for("profile"))

        if len(name) < 2:
            if is_ajax:
                return jsonify(success=False, error="Name too short.")
            flash("Name must be at least 2 characters long.", "danger")
            return redirect(url_for("profile"))

        if grade == "10-12":
            pass
        elif grade.isdigit() and 1 <= int(grade) <= 12:
            pass
        else:
            flash("Invalid grade.", "danger")
            return redirect(url_for("profile"))

        
        if not board:
            if is_ajax:
                return jsonify(success=False, error="Board is required.")
            flash("Board is required.", "danger")
            return redirect(url_for("profile"))

        if not age.isdigit() or int(age) < 1:
            if is_ajax:
                return jsonify(success=False, error="Invalid age.")
            flash("Please enter a valid age.", "danger")
            return redirect(url_for("profile"))

        dob = None
        if dob_str:
            try:
                dob = datetime.strptime(dob_str, "%Y-%m-%d").date()
                if dob > datetime.today().date():
                    raise ValueError
            except ValueError:
                if is_ajax:
                    return jsonify(success=False, error="Invalid Date of Birth.")
                flash("Invalid Date of Birth.", "danger")
                return redirect(url_for("profile"))
            
            # ---------- PHONE VALIDATION (CRITICAL FIX) ----------
            digits = re.sub(r"\D", "", phone_number)

            if not is_valid_indian_mobile(digits):
                if is_ajax:
                    return jsonify(success=False, error="Invalid mobile number.")
                flash("Enter a valid 10-digit Indian mobile number.", "danger")
                return redirect(url_for("profile"))

            full_phone = current_user.country_code + digits

            # ❗ Ensure phone number is unique (important for Google users)
            existing_user = User.query.filter(
                User.phone_number == full_phone,
                User.id != current_user.id
            ).first()

            if existing_user:
                if is_ajax:
                    return jsonify(success=False, error="Mobile number already in use.")
                flash("Mobile number already in use.", "danger")
                return redirect(url_for("profile"))

        # ---------- SAVE ----------
        current_user.name = name
        current_user.grade = grade
        current_user.board = board
        current_user.age = int(age)
        current_user.dob = dob
        current_user.city = city or None
        current_user.postal_code = postal_code or None
        current_user.timezone = timezone or None
        current_user.whatsapp_consent = 'whatsapp_consent' in request.form
        current_user.newsletter_consent = 'newsletter_consent' in request.form
        current_user.phone_number = full_phone   # ✅ use validated phone
        current_user.profile_completed = True

        try:
            db.session.commit()


            # decide redirect target
            if current_user.grade == "10-12" or current_user.board == "CBSE-ICSE":
                redirect_url = url_for("exam_combo_page")
            else:
                redirect_url = url_for("serve_index")

            if is_ajax:
                return jsonify(success=True, redirect=redirect_url)

            flash("Profile updated successfully!", "success")
            return redirect(redirect_url)

        except Exception as e:
            db.session.rollback()
            print("Profile update error:", e)

            if is_ajax:
                return jsonify(success=False, error="Server error.")

            flash("Something went wrong.", "danger")
            return redirect(url_for("profile"))

    # ---------- GET ----------
    return render_template("profile.html", errors={}, form={})


@app.route('/settings', methods=['GET', 'POST'])
@login_required
def settings():
    if not current_user.password_hash:
        flash('You logged in with Google and do not have a local password to change.', 'info')
        return render_template('settings.html')

    if request.method == 'POST':
        old_password = request.form.get('old_password')
        new_password1 = request.form.get('new_password1')
        new_password2 = request.form.get('new_password2')

        if not current_user.check_password(old_password):
            flash('Incorrect old password.', 'danger')
            return redirect(url_for('settings'))

        if new_password1 != new_password2:
            flash('New passwords do not match.', 'warning')
            return redirect(url_for('settings'))

        if len(new_password1) < 8:
            flash('Password requires at least 8 characters.', 'warning')
            return redirect(url_for('settings'))

        if not re.search(r'[A-Z]', new_password1):
            flash('Password requires at least one uppercase letter.', 'warning')
            return redirect(url_for('settings'))

        if not re.search(r'[a-z]', new_password1):
            flash('Password requires at least one lowercase letter.', 'warning')
            return redirect(url_for('settings'))

        if not re.search(r'\d', new_password1):
            flash('Password requires at least one digit.', 'warning')
            return redirect(url_for('settings'))

        if not re.search(r'[!@#$%^&*()_+=\-\[\]{};\'\\:"|,.<>\/?~`]', new_password1):
            flash('Password requires at least one special character.', 'warning')
            return redirect(url_for('settings'))

        if re.search(r'\s', new_password1):
            flash('Password cannot contain spaces.', 'warning')
            return redirect(url_for('settings'))

        if current_user.check_password(new_password1):
            flash('New password cannot be the same as the old password.', 'warning')
            return redirect(url_for('settings'))

        # ✅ SET PASSWORD
        current_user.set_password(new_password1)

        try:
            db.session.commit()
            flash('Password updated successfully!', 'success')

            # ✅ POPUP RESPONSE
            if request.headers.get("X-Requested-With") == "XMLHttpRequest":
                return jsonify({"success": True})

            return redirect(url_for('settings'))

        except Exception as e:
            db.session.rollback()
            flash(f'An error occurred: {e}', 'danger')
            print(f"Error updating password: {e}")
            return redirect(url_for('settings'))

    return render_template('settings.html')



@app.route("/whatsapp_webhook", methods=['POST'])
def whatsapp_webhook():
    """This route is called by Twilio when a user sends you a message."""
    
    # 1. Get the message a user sent
    incoming_message = request.form.get('Body', '').lower()
    
    reply_text = ""
    
    # --- START OF MODIFIED LOGIC ---

    if "course" in incoming_message:
        reply_text = "Hi! We offer courses in Math, Science, and English. Which one are you interested in?"
    
    elif "price" in incoming_message:
        reply_text = "Our courses start at $100. You can find more details on our website."
    
    elif "hello" in incoming_message:
        reply_text = "Hi there! I'm the MathGen AI assistant. How can I help you with our courses today?"
    
    # --- NEW PART 1: Handle "worksheet" query ---
    elif "worksheet" in incoming_message:
        # ** REPLACE with your actual website URL **
        reply_text = "You can generate worksheets by visiting our main app at https://your-website-url.com" 
    
    # --- NEW PART 2: Updated fallback message with contact info ---
    else:
        # ** REPLACE with your real email and phone **
        reply_text = (
            "Sorry, I don't understand that query. You can ask me about 'courses' or 'prices'. \n\n"
            "For more help, please contact us at:\n"
            "📧 support@mathgen.com\n"
            "📞 +91 99999 88888"
        )
        
    # --- END OF MODIFIED LOGIC ---

    # 3. Create a reply and send it back to Twilio
    resp = MessagingResponse()
    resp.message(reply_text)
    
    return str(resp)



@app.route("/test/<int:attempt_id>")
@login_required
def take_test(attempt_id):
    attempt = MockAttempt.query.get_or_404(attempt_id)
    
    # 🔒 SECURITY CHECK: Does this attempt belong to the logged-in user?
    if attempt.user_id != current_user.id:
        flash("Unauthorized access!", "danger")
        return redirect(url_for("serve_index"))

    test = MockTest.query.get_or_404(attempt.test_id)
    questions = MockQuestion.query.filter_by(test_id=test.id).all()

    return render_template("test_page.html", test=test, attempt=attempt, questions=questions)

@app.route("/review/<int:attempt_id>")
@login_required
def review_attempt(attempt_id):
    attempt = MockAttempt.query.get_or_404(attempt_id)

    if attempt.user_id != current_user.id:
        flash("Not allowed", "danger")
        return redirect(url_for("my_scores"))

    test = MockTest.query.get(attempt.test_id)
    questions = MockQuestion.query.filter_by(test_id=test.id).order_by(MockQuestion.qno).all()

    answers = {a.question_id: a.selected_option for a in attempt.answers}

    rows = []
    for q in questions:
        selected = answers.get(q.id)
        rows.append({
    "qno": q.qno,
    "question": clean_ai_text(q.question_text),
    "options": [clean_ai_text(o) for o in q.options()],
    "correct": clean_ai_text(q.correct_option),
    "selected": selected,
    "is_correct": selected == q.correct_option_index,
    "explanation": clean_ai_text(q.explanation)
})


    return render_template("review.html", rows=rows, attempt=attempt, test=test)


def format_answers_numbered(text):
    """
    Forces each numbered answer onto a new line.
    Converts: '1) ... 2) ... 3) ...'
    Into:
    1) ...
    2) ...
    3) ...
    """
    if not text:
        return ""

    # Normalize spaces
    text = re.sub(r"\s+", " ", text)

    # Force line break before each number like "1)" or "2)"
    text = re.sub(r"(\d+)\)", r"\n\1)", text)

    return text.strip()


def format_questions_for_exam(text):
    lines = []
    q_no = 1

    for block in text.split("\n"):
        block = block.strip()
        if not block:
            continue

        # Detect question lines
        if block[0].isdigit() or block.lower().startswith("q"):
            lines.append(f"\nQ{q_no}. {block}\n")
            lines.append("________________________________________\n")
            lines.append("________________________________________\n")
            lines.append("________________________________________\n")
            q_no += 1
        else:
            lines.append(block + "\n")

    return "".join(lines)

def normalize_job_questions(questions_json):
    output = []
    current_section = None
    qno = 1

    for q in questions_json:
        section = q.get("section", "General")
        year = q.get("year", "Previous Year")
        question = clean_ai_text(q.get("question", ""))
        options = q.get("options", [])

        if section != current_section:
            output.append(f"\n=== {section.upper()} ===\n")
            current_section = section

        output.append(f"{qno}. {question} ({year})")
        for idx, opt in enumerate(options):
            output.append(f"   {chr(65+idx)}. {clean_ai_text(opt)}")

        output.append("")  # blank line
        qno += 1

    return "\n".join(output)



def get_output_format():
    worksheet_type = request.form.get("worksheet_type")

    if worksheet_type == "job":
        return (request.form.get("job_format") or "pdf").lower()

    if worksheet_type == "school":
        return (request.form.get("school_format") or "pdf").lower()

    return "pdf"

@app.route("/exam-combo")
@login_required
def exam_combo_page():
    return render_template(
        "exam_combo.html",
        user_grade=current_user.grade,
        user_board=current_user.board
    )


@app.route("/generate-exam-combo", methods=["POST"])
@login_required
def generate_exam_combo():
    return handle_exam_combo(request)

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import zipfile
from flask import send_file
from google import genai

os.makedirs(TEMP_DIR, exist_ok=True)


def extract_json_from_ai(text, expect="question"):
    if not text:
        return None

    text = text.strip()

    match = re.search(r"\[\s*\{[\s\S]*?\}\s*\]", text)
    if not match:
        return None

    try:
        data = json.loads(match.group())
        if not isinstance(data, list):
            return None

        # 🔒 HARD VALIDATION
        for item in data:
            if expect == "question" and "question" not in item:
                return None
            if expect == "answer" and not any(k in item for k in ("answer", "solution")):
                return None

        return data

    except Exception:
        return None

  
def clean_ai_text(text):
    # Convert LaTeX fractions to a/b
    text = re.sub(r"\\frac\s*\{([^}]+)\}\{([^}]+)\}", r"\1/\2", text)

    if not text:
        return ""

    # ---------- LaTeX → Unicode Math ----------
    replacements = {
        # Greek
        "\\alpha": "α", "\\beta": "β", "\\gamma": "γ",
        "\\delta": "δ", "\\theta": "θ", "\\pi": "π",
        "\\lambda": "λ", "\\mu": "μ", "\\sigma": "σ",

        # Roots & powers
        "\\sqrt": "√",
        "^2": "²", "^3": "³",
        "^4": "⁴", "^5": "⁵",

        # Operators
        "\\times": "×",
        "\\cdot": "·",
        "\\pm": "±",
        "\\div": "÷",

        # Relations
        "\\le": "≤",
        "\\ge": "≥",
        "\\neq": "≠",
        "\\approx": "≈",

        # Brackets
        "\\left(": "(",
        "\\right)": ")",
        "\\left[": "[",
        "\\right]": "]",
        "\\left\\{": "{",
        "\\right\\}": "}",

        # Noise
        "$": "",
        "\\(": "",
        "\\)": "",
        "\\,": " ",
        "\\;": " ",
        "\\!": "",
        "\\": "",
        
        "tan^{-1}": "tan⁻¹",
        "sin^{-1}": "sin⁻¹",
        "cos^{-1}": "cos⁻¹",
        "{": "",
        "}": "",

    }

    for k, v in replacements.items():
        text = text.replace(k, v)

    # ---------- Fix remaining math formatting ----------
    # sqrt(x) → √x
    text = re.sub(r"√\(([^)]+)\)", r"√\1", text)

    # Remove double spaces
    text = re.sub(r"\s{2,}", " ", text)

    # Clean lines
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return "\n".join(lines)
  
def normalize_questions(questions_json):
    lines = []

    for i, q in enumerate(questions_json, 1):
        question = clean_ai_text(q.get("question", "")).strip()

        # 🚨 BLOCK answers sneaking in
        if re.match(r"^[-+]?\d|^\(.*\)$", question):
            raise ValueError("Answer detected where question expected")

        lines.append(f"{i}) {question}")
        lines.append("")

    return "\n".join(lines)



def normalize_answers(text):
    """
    Forces STRICT competitive-exam answer format:
    1. (a) Option text
    """

    if not text:
        raise ValueError("Empty answer response")

    text = clean_ai_text(text)

    lines = []
    for line in text.splitlines():
        line = line.strip()

        # Accept only formats like: 1. (a) XYZ
        if re.match(r"^\d+\.\s*\([a-d]\)\s+.+", line, re.I):
            lines.append(line)

    if not lines:
        raise ValueError("No valid answers detected")

    return "\n".join(lines)

# --- PDF Generation Class ---
class CustomPDF(FPDF):
    def __init__(self, title, sub_info, header_text="", footer_text=""):
        super().__init__()
        FONT_PATH = os.path.join(BASE_DIR, "fonts", "DejaVuSans.ttf")
        self.add_font("DejaVu", "", FONT_PATH, uni=True)
        self.worksheet_title = title
        self.sub_info = sub_info
        self.header_text = header_text
        self.footer_text = footer_text

    def header(self):
        self.set_font("DejaVu", "", 10)
        self.cell(0, 8, self.header_text, 0, 1, "C")

        self.set_font("DejaVu", "", 16)
        self.cell(0, 10, self.worksheet_title, 0, 1, "C")

        self.set_font("DejaVu", "", 10)
        sub_header_text = (
            f"Date: {self.sub_info['date']}   |   "
            f"Marks: {self.sub_info['marks']}   |   "
            f"Topic: {self.sub_info['sub-title']}"
        )
        self.cell(0, 8, sub_header_text, 0, 1, "C")
        self.ln(6)


    def footer(self):
        self.set_y(-15)
        self.set_font("DejaVu", "", 8)

        # --- Footer Left Text ---
        self.cell(0, 8, self.footer_text, 0, 0, "L")

        # --- Page Number (Right) ---
        self.cell(0, 8, f"Page {self.page_no()}", 0, 0, "R")

# --- File Creation Functions ---
def create_pdf(content, title, sub_title_info, filename=None):
    if not filename:
        filename = "Worksheet.pdf"

    file_path = os.path.join(TEMP_DIR, filename)

    pdf = CustomPDF(
        title,
        sub_title_info,
        header_text=GLOBAL_HEADER,
        footer_text=GLOBAL_FOOTER
    )

    pdf.add_page()
    pdf.set_font("DejaVu", "", 12)
    
    # Calculate usable width (Page width - left margin - right margin)
    usable_width = pdf.w - pdf.l_margin - pdf.r_margin

    # SPLIT ONLY BY NEWLINES: Do not strip all whitespace or you lose the math structure
    for line in content.splitlines():
        line = line.strip()
        if not line:
            pdf.ln(5)  # Proper vertical space for empty lines
            continue

        # This ensures the question/answer stays on one line unless it truly exceeds page width
        pdf.multi_cell(usable_width, 8, line, border=0, align="L")
        pdf.ln(2) # Small gap between different answers

    pdf.output(file_path)
    return file_path



import io
import zipfile
from datetime import datetime
from google import genai
from flask import send_file
import csv


def get_combo_syllabus(board, grade):
    TOPICS_CSV = os.path.join(BASE_DIR, "topics.csv")

    grades = {"10", "11", "12"} if grade == "10-12" else {grade}
    boards = {"cbse", "icse"} if board.lower() == "cbse-icse" else {board.lower()}

    topics = []

    with open(TOPICS_CSV, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            if (
                row["board"].lower() in boards
                and row["grade"] in grades
                and row["subject"].lower() == "mathematics"
            ):
                major = row["major_topic"].strip()
                minor = row["minor_topic"].strip()

                topics.append(f"{major} - {minor}" if minor else major)

    return list(set(topics))


def handle_exam_combo(request):
    grade = request.form.get("grade")
    board = request.form.get("board")
    paper_type = request.form.get("paper_type")
    year = request.form.get("year")
    include_solutions = request.form.get("include_solutions") == "1"
    include_detailed = request.form.get("include_detailed_solutions") == "1"
    if include_detailed:
        include_solutions = True


    if not grade or not board or not paper_type:
        raise ValueError("Missing required fields")

    syllabus_topics = get_combo_syllabus(board, grade)
    if not syllabus_topics:
        raise ValueError("No syllabus found")

    syllabus_text = ", ".join(syllabus_topics)

    client = genai.Client(api_key=os.environ.get("GENAI_API_KEY"))

    if paper_type == "past":
        prompt = f"""
Generate a REALISTIC past exam paper.

Year: {year}

Syllabus (MUST MIX QUESTIONS FROM ALL):
{syllabus_text}

Rules:
- Exam-level
- NO answers
- EXACTLY 15 questions
- Plain text only

Return STRICT JSON:
[{{ "question": "..." }}]
"""
        title = f"Grade {grade} {board} Past Paper ({year})"
    else:
        prompt = f"""
Generate a FULL SYLLABUS mock exam paper.

Syllabus (MUST MIX QUESTIONS FROM ALL):
{syllabus_text}

Rules:
- Exam-level
- NO answers
- EXACTLY 15 questions
- Plain text only

Return STRICT JSON:
[{{ "question": "..." }}]
"""
        title = f"Grade {grade} {board} Mock Paper"

    response = client.models.generate_content(
        model="models/gemini-flash-latest",
        contents=prompt
    )

    questions_json = extract_json_from_ai(response.text)
    if not isinstance(questions_json, list):
        raise ValueError("AI did not return a JSON list")


    # 🔒 HARD VALIDATION: QUESTIONS ONLY
    if not questions_json or not isinstance(questions_json, list):
        raise ValueError("AI did not return valid questions")

    for q in questions_json:
        if "question" not in q:
            raise ValueError("AI returned answers instead of questions")

    worksheet_text = normalize_questions(questions_json)
    # 🚨 HARD BLOCK — exam paper must NOT contain answers
    if re.search(r"^\d+[\)\.]\s*[-+]?\d", worksheet_text, re.M):
        raise ValueError("CRITICAL: Answers detected in EXAM worksheet")



    solution_text = None
    detailed_solution_text = None

    if include_solutions and not include_detailed:
        answer_prompt = f"""
Return ONLY the final answers.

STRICT RULES:
- NO questions
- NO explanations
- One answer per line
- Format EXACTLY:

1. (a) Option text
2. (b) Option text

Questions:
{worksheet_text}
"""
        resp = client.models.generate_content(
            model="models/gemini-flash-latest",
            contents=answer_prompt
        )
        solution_text = normalize_answers(clean_ai_text(resp.text))

    elif include_detailed:
        detailed_prompt = f"""
You are an expert examiner.

Generate DETAILED solutions.

RULES:
- Question by question
- Step-by-step explanation
- Simple student language
- Plain text only
- NO markdown

FORMAT:

Q1.
Final Answer: (a) Option
Explanation:
Step 1: ...
Step 2: ...
Conclusion: ...

Questions:
{worksheet_text}
"""
        resp = client.models.generate_content(
            model="models/gemini-flash-latest",
            contents=detailed_prompt
        )
        detailed_solution_text = clean_ai_text(resp.text)

        a_response = client.models.generate_content(
            model="models/gemini-flash-latest",
            contents=answer_prompt
        )

        try:
            solution_text = normalize_answers(clean_ai_text(a_response.text))
        except ValueError as e:
            return jsonify({"error": str(e)}), 500


        # 🚨 HARD SAFETY CHECKS (ADD HERE)
        if solution_text.strip() == worksheet_text.strip():
            raise ValueError("AI returned questions instead of answers")


    info = {
        "date": datetime.now().strftime("%d %b %Y"),
        "marks": "___ / 80",
        "sub-title": "Full Syllabus"
    }

    files = []
    ws = create_pdf(
        worksheet_text,     # ✅ DIRECT QUESTIONS
        title,
        info,
        filename="Worksheet.pdf"
    )


    files.append(("Worksheet.pdf", ws))

    if solution_text:
        sol = create_pdf(
            solution_text,
            f"{title} - Solutions",
            info,
            filename="Answer_Sheet.pdf"
        )
        files.append(("Answer_Sheet.pdf", sol))
        
    if detailed_solution_text:
        detailed_path = create_pdf(
            detailed_solution_text,
            f"{title} - Detailed Solutions",
            info,
            filename="Detailed_Solutions.pdf"
        )
        files.append(("Detailed_Solutions.pdf", detailed_path))




    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w") as zipf:
        for name, path in files:
            zipf.write(path, name)

    zip_buffer.seek(0)
    return send_file(zip_buffer, as_attachment=True, download_name="Exam_Papers.zip")

def clear_temp_dir():
    for f in os.listdir(TEMP_DIR):
        try:
            os.remove(os.path.join(TEMP_DIR, f))
        except:
            pass
        
# ✅ AUTO CLEAN TEMP FILES AFTER EVERY RESPONSE
@app.after_request
def cleanup_temp(response):
    if response.mimetype == "application/zip":
        clear_temp_dir()
    return response


def normalize_exam_key(authority):
    authority = authority.upper()
    if "SSC" in authority:
        return "SSC CGL"
    if "IBPS PO" in authority:
        return "IBPS PO"
    if "SBI PO" in authority:
        return "SBI PO"
    if "RRB" in authority:
        return "RRB NTPC"
    if "NDA" in authority:
        return "UPSC NDA"
    if "GATE" in authority:
        return "GATE"
    return authority



@app.route('/generate-worksheet', methods=['GET', 'POST'])
@login_required
def generate_worksheet():
    if request.form.get("grade") == "10-12" or request.form.get("board") == "CBSE-ICSE":
        return jsonify({
            "error": "Combo syllabus must be generated via Exam Papers section."
            
        }), 400

    if request.method == 'GET':
        return render_template('index.html')

    title = "Math Worksheet"
    info = {
        "date": datetime.now().strftime("%d %b %Y"),
        "time": datetime.now().strftime("%I:%M %p"),
        "marks": "___ / 50",
        "sub-title": "General"
    }

    try:
        client = genai.Client(api_key=os.environ.get("GENAI_API_KEY"))

        worksheet_type = request.form.get("worksheet_type")
        if worksheet_type not in ("school", "job"):
            return jsonify({"error": "Invalid worksheet type"}), 400


        job_type = request.form.get("job_type") or None
        exam_type = request.form.get("exam_type")
        exam_authority = request.form.get("exam_authority")
        
        # =====================================================
        # JOB / COMPETITIVE WORKSHEET GENERATION
        # =====================================================
        if worksheet_type == "job":
            if not job_type or not exam_type or not exam_authority:
                flash("Please select job category, exam type and authority.", "warning")
                return redirect(url_for("jobseekers"))

        
            output_format = get_output_format()
            include_answers = 'answer_key' in request.form
            job_type_upper = job_type.upper() if job_type else "JOB"
            
            exam_key = normalize_exam_key(exam_authority)
            exam_structure = EXAM_STRUCTURES.get(exam_key)

            if not exam_structure:
                return jsonify({
                    "error": f"No predefined structure found for {exam_key}"
                }), 400


            sections_text = ""
            total_questions = 0

            for section, count in exam_structure.items():
                sections_text += f"- {section}: {count} questions\n"
                total_questions += count


            job_prompt = f"""
You are a REAL competitive exam paper setter.

Your task is to generate a paper that is INDISTINGUISHABLE from an actual official competitive exam paper.

🚨 ABSOLUTE NON-NEGOTIABLE RULES 🚨
If ANY rule is violated, the response is INVALID.

1️⃣ EXAM AUTHENTICITY (MANDATORY)
- Questions MUST match REAL past competitive exam papers
- Style, depth, framing, and difficulty must be IDENTICAL to official exams
- Questions must feel like they came from an actual exam hall

2️⃣ STRICT SYLLABUS BOUNDARY
- Generate questions ONLY from the OFFICIAL SYLLABUS of the selected exam
- DO NOT include:
  ❌ college/university theory questions
  ❌ GATE-style academic questions (unless the exam IS GATE)
  ❌ school-level or generic CS questions
  ❌ definition-only or textbook recall questions

3️⃣ QUESTION STYLE CONTROL
- Questions MUST be:
  ✔ scenario-based
  ✔ application-oriented
  ✔ decision-based
  ✔ exam-contextual (banking / government / engineering context as applicable)

- Questions MUST NOT be:
  ❌ “What is X?”
  ❌ “Which data structure is used for Y?”
  ❌ pure theory recall
  ❌ memorization-only MCQs

4️⃣ EXAM-SPECIFIC INTELLIGENCE (CRITICAL)
- Understand that EACH competitive exam has a UNIQUE nature
- DO NOT reuse question styles from other exams
  (e.g., DO NOT mix GATE-style questions into IBPS/SSC)
- Adapt question framing EXACTLY to the selected exam

5️⃣ YEAR HANDLING (IMPORTANT)
- You may reference ANY year (2010–Present)
- Do NOT fixate on a single year
- Year is for realism ONLY, not restriction

6️⃣ CONTENT SAFETY
- NO answers inside questions
- NO explanations
- NO hints
- NO formatting tricks
- NO markdown
- NO LaTeX

7️⃣ FORMAT CONTROL (STRICT)
Return ONLY valid JSON.
NO extra text.
NO commentary.
NO markdown.

FORMAT:
[
  {
    "section": "Exact exam section name",
    "year": "Previous Year",
    "question": "Realistic competitive exam question",
    "options": ["A", "B", "C", "D"]
  }
]

If you are unsure whether a question truly belongs to the exam,
DO NOT generate it.
"""



            job_response = client.models.generate_content(
                model="models/gemini-flash-latest",
                contents=job_prompt
            )

            questions_json = extract_json_from_ai(job_response.text)
            if not isinstance(questions_json, list):
                raise ValueError("AI did not return a JSON list")


            if not questions_json or not isinstance(questions_json, list):
                print("❌ RAW JOB RESPONSE ↓↓↓")
                print(job_response.text)
                return jsonify({"error": "AI failed to generate job questions"}), 500

            worksheet_questions_text = normalize_job_questions(questions_json)
            # 🚨 HARD BLOCK — worksheet must NOT contain answers
            if re.search(r"^\d+[\)\.]\s*[-+]?\d", worksheet_questions_text, re.M):
                raise ValueError("CRITICAL: Answers detected in JOB worksheet content")


            title = f"{job_type_upper} {exam_type.upper()} Worksheet"
            info["sub-title"] = exam_authority

            solution_answers_text = None

            if include_answers:
                answers_prompt = f"""
You are generating ONLY the final answer key for a competitive exam.

STRICT RULES (NON-NEGOTIABLE):
- DO NOT repeat questions
- DO NOT explain
- DO NOT add extra text
- EACH answer must be in THIS EXACT FORMAT:

1. (a) Option text
2. (c) Option text
3. (b) Option text

Rules:
- Use a, b, c, or d ONLY
- Option text must match the correct option
- One answer per line ONLY

Questions:
{worksheet_questions_text}
"""

                a_response = client.models.generate_content(
                    model="models/gemini-flash-latest",
                    contents=answers_prompt
                )

                try:
                    solution_answers_text = normalize_answers(clean_ai_text(a_response.text))
                except ValueError as e:
                    return jsonify({"error": str(e)}), 500

                
                # 🚨 HARD SAFETY CHECK — STOP ANSWERS LEAKING INTO WORKSHEET
                # 🚨 HARD SAFETY CHECK — STOP ANSWERS LEAKING INTO WORKSHEET
                if solution_answers_text:
                    common = set(solution_answers_text.splitlines()) & set(worksheet_questions_text.splitlines())
                    if len(common) > 2:
                        raise ValueError(
                            "CRITICAL: Answers leaked into worksheet (content overlap detected)"
                        )



            if solution_answers_text and (
                solution_answers_text.strip() == worksheet_questions_text.strip()
            ):
                raise ValueError("AI returned questions instead of answers")






            files_to_zip = []

            # ---------- IMAGE ----------
            if output_format in ["jpg", "png", "gif"]:
                ws = create_image(
                    worksheet_questions_text,
                    title,
                    info,
                    filename=f"Worksheet.{output_format}"
,
                    fmt=output_format
                )
                files_to_zip.append((f"Worksheet.{output_format}", ws))

                if include_answers and solution_answers_text:
                    sol = create_image(
                        solution_answers_text,
                        f"{title} - Solutions",
                        info,
                        filename=f"Solutions.{output_format}",
                        fmt=output_format
                    )
                    files_to_zip.append((f"Solutions.{output_format}", sol))




            # ---------- TXT ----------
            elif output_format == "txt":
                ws = create_txt(
                    worksheet_questions_text,
                    title,
                    info,
                    f"Worksheet.txt"
                )
                files_to_zip.append(("Worksheet.txt", ws))

                if include_answers and solution_answers_text:
                    sol = create_txt(
                        solution_answers_text,
                        f"{title} - Solutions",
                        info,
                        f"Solutions.txt"
                    )
                    files_to_zip.append(("Solutions.txt", sol))



            # ---------- DOCX ----------
            elif output_format == "docx":
                ws = create_docx(
                    worksheet_questions_text,
                    title,
                    info,
                    f"Worksheet.docx"
                )
                files_to_zip.append(("Worksheet.docx", ws))

                if include_answers and solution_answers_text:
                    sol = create_docx(
                        solution_answers_text,
                        f"{title} - Solutions",
                        info,
                        f"Solutions.docx"
                    )
                    files_to_zip.append(("Solutions.docx", sol))




            # ---------- PDF ----------
            elif output_format == "pdf":
                worksheet_path = create_pdf(
                    worksheet_questions_text,   # ✅ DIRECT QUESTIONS
                    title,
                    info,
                    filename="Worksheet.pdf"
                )

                files_to_zip.append(("Worksheet.pdf", worksheet_path))


                if include_answers and solution_answers_text:
                    solution_path = create_pdf(
                        solution_answers_text,
                        f"{title} - Solutions",
                        info,
                        filename="Answer_Sheet.pdf"
                    )
                    files_to_zip.append(("Answer_Sheet.pdf", solution_path))



            else:
                return jsonify({"error": "Invalid format selected"}), 400
            
            if not files_to_zip:
                return jsonify({"error": "Nothing to download"}), 400


            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
                for name, path in files_to_zip:
                    with open(path, "rb") as f:
                        zipf.writestr(name, f.read())

            zip_buffer.seek(0)
            
            return send_file(
    zip_buffer,
    mimetype="application/zip",
    as_attachment=True,
    download_name="Job_Worksheet.zip"
)





        # =====================================================
        # FILE UPLOAD FLOW (ANSWER KEY ONLY)
        # =====================================================
        if (
            'worksheet_file' in request.files
            and request.files['worksheet_file']
            and request.files['worksheet_file'].filename.strip() != ""
):


            file = request.files['worksheet_file']
            filename = file.filename.lower()
            file.stream.seek(0)  # 🔥 CRITICAL FIX

            # ---------- READ CONTENT FROM UPLOADED FILE ----------
            if filename.endswith(".pdf"):
                worksheet_questions_text = get_text_from_pdf(file)

            elif filename.endswith(".docx"):
                worksheet_questions_text = get_text_from_docx(file)

            elif filename.endswith(".txt"):
                worksheet_questions_text = file.read().decode("utf-8", errors="ignore")

            elif filename.endswith((".png", ".jpg", ".jpeg", ".gif")):
                worksheet_questions_text = get_text_from_image(file)

            else:
                return jsonify({"error": "Unsupported uploaded file format"}), 400

            include_answers = 'answer_key' in request.form


            # ✅ FORCE output format SAME AS uploaded file
            if filename.endswith(".pdf"):
                output_format = "pdf"
            elif filename.endswith(".docx"):
                output_format = "docx"
            elif filename.endswith(".txt"):
                output_format = "txt"
            elif filename.endswith(".png"):
                output_format = "png"
            elif filename.endswith(".jpg") or filename.endswith(".jpeg"):
                output_format = "jpg"
            elif filename.endswith(".gif"):
                output_format = "gif"
            else:
                return jsonify({"error": "Unsupported uploaded file format"}), 400

            if not worksheet_questions_text or not worksheet_questions_text.strip():
                return jsonify({"error": "Could not read worksheet"}), 400


            prompt = f"""
You are generating ONLY the final answer key.

STRICT RULES:
- NO questions
- NO explanations
- EXACT format ONLY:

1. (a) Option text
2. (b) Option text
3. (c) Option text

Rules:
- Use a, b, c, or d ONLY
- One answer per line

Worksheet:
{worksheet_questions_text}
"""


            response = client.models.generate_content(
                model="models/gemini-flash-latest",
                contents=prompt
            )

            try:
                solution_answers_text = normalize_answers(clean_ai_text(response.text))
            except ValueError as e:
                return jsonify({"error": str(e)}), 500

            # 🚨 FINAL SAFETY CHECK — answers must NOT contain questions
            if "?" in solution_answers_text:
                raise ValueError("AI returned questions instead of answers")


            files_to_zip = []

            # ========= IMAGE =========
            # ---------- IMAGE FORMATS ----------
            if output_format in ["jpg", "png", "gif"]:
                ws = create_image(
                    worksheet_questions_text,
                    title,
                    info,
                    filename=f"Worksheet.{output_format}"
,
                    fmt=output_format
                )
                files_to_zip.append((f"Worksheet.{output_format}", ws))


                if include_answers and solution_answers_text:
                    sol = create_image(
                        solution_answers_text,
                        f"{title} - Solutions",
                        info,
                        filename=f"Solutions.{output_format}",
                        fmt=output_format
                    )
                    files_to_zip.append((f"Solutions.{output_format}", sol))





            # ---------- TXT ----------
            elif output_format == "txt":
                worksheet_path = create_txt(
                    worksheet_questions_text,
                    title,
                    info,
                    filename=f"Worksheet.txt"
                )
                files_to_zip.append(("Worksheet.txt", worksheet_path))

                if include_answers and solution_answers_text:
                    solution_path = create_txt(
                        solution_answers_text,
                        f"{title} - Solutions",
                        info,
                        filename=f"Solutions.txt"
                    )
                    files_to_zip.append(("Solutions.txt", solution_path))


            # ---------- DOCX ----------
            elif output_format == "docx":

                worksheet_path = create_docx(
                    worksheet_questions_text,
                    title,
                    info,
                    filename= "Worksheet.docx"
                )
                files_to_zip.append(("Worksheet.docx", worksheet_path))

                if include_answers and solution_answers_text:

                    solution_path = create_docx(
                        solution_answers_text,
                        f"{title} - Solutions",
                        info,
                        filename="Solutions.docx"
                    )
                    files_to_zip.append(("Solutions.docx", solution_path))


            # ---------- PDF ----------
            elif output_format == "pdf":
                worksheet_path = create_pdf(
                    worksheet_questions_text,   # ✅ DIRECT QUESTIONS
                    title,
                    info,
                    filename="Worksheet.pdf"
                )

                files_to_zip.append(("Worksheet.pdf", worksheet_path))


                if include_answers and solution_answers_text:
                    solution_path = create_pdf(
                        solution_answers_text,
                        f"{title} - Solutions",
                        info
                    )
                    files_to_zip.append(
                        ("Solutions.pdf", solution_path))
                    



            else:
                return jsonify({"error": "Invalid format selected"}), 400
            
            if not files_to_zip:
                return jsonify({"error": "Nothing to download"}), 400


                
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
                for zip_name, file_path in files_to_zip:
                    with open(file_path, "rb") as f:
                        zipf.writestr(zip_name, f.read())

            zip_buffer.seek(0)

            return send_file(
                zip_buffer,
                mimetype="application/zip",
                as_attachment=True,
                download_name="Worksheet_Solutions.zip"
            )
            
        # =====================================================
        # NEW WORKSHEET GENERATION FLOW
        # =====================================================
        if worksheet_type == "school":
            topic = request.form.get('topic')
            if not topic:
                return jsonify({"error": "Topic is required for school worksheet"}), 400

            grade = request.form.get('grade') or current_user.grade
            board = request.form.get('board')
            subtopic = request.form.get('subtopic') or "General"
            difficulty = request.form.get('difficulty', 'Easy')
            output_format = get_output_format()
            include_answers = 'answer_key' in request.form


            # 🔒 Subtraction digit control
            digit_rule = ""
            if topic and any(op in topic.lower() for op in ["addition", "subtraction", "multiplication", "division"]):
                if "2" in subtopic:
                    digit_rule = (
                        "STRICT RULE: Use ONLY 2-digit numbers (10 to 99). "
                        "DO NOT use single-digit numbers."
                    )
                elif "3" in subtopic:
                    digit_rule = (
                        "STRICT RULE: Use ONLY 3-digit numbers (100 to 999). "
                        "DO NOT use smaller numbers."
                    )
                elif "4" in subtopic:
                    digit_rule = (
                        "STRICT RULE: Use ONLY 4-digit numbers (1000 to 9999). "
                        "DO NOT use smaller numbers."
                    )


            questions_prompt = f"""
You are a STRICT academic question generator.

Your task is to generate questions ONLY from the given topic and subtopic.

🚨 ABSOLUTE RULES (NON-NEGOTIABLE):
- Generate EXACTLY 15 questions
- Questions MUST belong strictly to:
  • Topic = "{topic}"
  • Subtopic = "{subtopic}"
- Questions MUST match:
  • Grade = {grade}
  • Board = {board}
  • Difficulty = {difficulty}
- DO NOT generate questions from any other topic or subtopic
- DO NOT include concepts from earlier/lower classes
- DO NOT include generic arithmetic unless explicitly part of the topic
- DO NOT reuse or paraphrase previously generated questions
- DO NOT include answers, hints, explanations, or examples
- Use ONLY plain English (no LaTeX, no markdown)

❌ INVALID EXAMPLES (DO NOT DO THIS):
- Mixing topics (e.g., addition inside decimals)
- Using easier-grade concepts
- Reusing previously generated worksheet patterns
- Generating unrelated arithmetic

✅ VALID OUTPUT FORMAT (STRICT JSON ONLY):

[
  {{
    "question": "A clear, student-ready question strictly from the given topic and subtopic",
    "answer_space_lines": 3
  }}
]

CONTEXT (MUST BE FOLLOWED EXACTLY):
Grade: {grade}
Board: {board}
Topic: {topic}
Subtopic: {subtopic}
Difficulty: {difficulty}

If ANY rule is violated, the response is INVALID.
"""

            q_response = client.models.generate_content(
                model="models/gemini-flash-latest",
                contents=questions_prompt
            )

            questions_json = extract_json_from_ai(q_response.text)
            if not isinstance(questions_json, list):
                raise ValueError("AI did not return a JSON list")


            if not questions_json or not isinstance(questions_json, list):
                print("❌ RAW GEMINI RESPONSE ↓↓↓")
                print(q_response.text)
                return jsonify({
                    "error": "AI returned invalid question format. Please try again."
                }), 500
            
            worksheet_questions_text = normalize_questions(questions_json)
            # 🚨 HARD BLOCK — worksheet must NOT contain answers
            if re.search(r"^\d+[\)\.]\s*[-+]?\d", worksheet_questions_text, re.M):
                raise ValueError("CRITICAL: Answers detected in SCHOOL worksheet content")

            
            # 🚨 HARD VALIDATION — STOP WRONG TOPIC GENERATION
            if "addition" in worksheet_questions_text.lower() and "decimal" not in worksheet_questions_text.lower():
                raise ValueError("AI generated questions outside the selected topic (Decimals)")



            solution_answers_text = None

            title = f"Grade {grade} Math Worksheet"
            info["sub-title"] = subtopic

            # -------- ANSWERS ----------
            if include_answers:
                answers_prompt = f"""
You are generating ONLY an answer key.

STRICT RULES (MANDATORY):
- ONLY answers
- NO explanations
- NO questions
- ONE answer per line
- Each answer MUST start on a new line
- FORMAT MUST BE EXACTLY:

1) Answer
2) Answer
3) Answer

If you violate format, response is INVALID.

Questions:
{worksheet_questions_text}
"""
                a_response = client.models.generate_content(
                    model="models/gemini-flash-latest",
                    contents=answers_prompt
                )
                try:
                    solution_answers_text = clean_ai_text(a_response.text)
                except ValueError as e:
                    return jsonify({"error": str(e)}), 500

                if solution_answers_text and (
                    worksheet_questions_text.strip() == solution_answers_text.strip()
                    ):
                    raise ValueError(
                        "Worksheet and Solutions content are identical. Aborting ZIP generation."
                        )




            # =====================================================
            # OUTPUT
            # =====================================================
            # =====================================================
            # ZIP OUTPUT (WORKSHEET + SOLUTIONS)
            # =====================================================
            # =====================================================
            # CREATE PDF FILES
            # =====================================================

            files_to_zip = []
            
            # ---------- IMAGE FORMATS ----------
            if output_format in ["jpg", "png", "gif"]:
                ws = create_image(
                    worksheet_questions_text,
                    title,
                    info,
                    filename=f"Worksheet.{output_format}"
,
                    fmt=output_format
                )
                files_to_zip.append((f"Worksheet.{output_format}", ws))


                if include_answers and solution_answers_text:
                    sol = create_image(
                        solution_answers_text,
                        f"{title} - Solutions",
                        info,
                        filename=f"Solutions.{output_format}",
                        fmt=output_format
                    )
                    files_to_zip.append((f"Solutions.{output_format}", sol))




            # ---------- TXT ----------
            elif output_format == "txt":
                worksheet_path = create_txt(
                    worksheet_questions_text,
                    title,
                    info,
                    filename=f"Worksheet.txt"
                )
                files_to_zip.append(("Worksheet.txt", worksheet_path))
            

                if include_answers and solution_answers_text:
                    solution_path = create_txt(
                        solution_answers_text,
                        f"{title} - Solutions",
                        info,
                        filename=f"Solutions.txt"
                    )
                    files_to_zip.append(("Solutions.txt", solution_path))
                    


            # ---------- DOCX ----------
            elif output_format == "docx":
                worksheet_path = create_docx(
                    worksheet_questions_text,
                    title,
                    info,
                    filename=f"Worksheet.docx"
                )
                files_to_zip.append(("Worksheet.docx", worksheet_path))
            

                if include_answers and solution_answers_text:
                    solution_path = create_docx(
                        solution_answers_text,
                        f"{title} - Solutions",
                        info,
                        filename=f"Solutions.docx"
                    )
                    files_to_zip.append(("Solutions.docx", solution_path))
                


            # ---------- PDF ----------
            elif output_format == "pdf":
                worksheet_path = create_pdf(
                    worksheet_questions_text,   # ✅ DIRECT QUESTIONS
                    title,
                    info,
                    filename="Worksheet.pdf"
                )

                files_to_zip.append(("Worksheet.pdf", worksheet_path))


                if include_answers and solution_answers_text:
                    solution_path = create_pdf(
                        solution_answers_text,
                        f"{title} - Solutions",
                        info
                    )
                    files_to_zip.append(("Solutions.pdf", solution_path))


            else:
                return jsonify({"error": "Invalid format selected"}), 400
            
            if not files_to_zip:
                return jsonify({"error": "Nothing to download"}), 400

            
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
                for zip_name, file_path in files_to_zip:
                    with open(file_path, "rb") as f:
                        zipf.writestr(zip_name, f.read())

            zip_buffer.seek(0)

            return send_file(
                zip_buffer,
                mimetype="application/zip",
                as_attachment=True,
                download_name="mathgen_worksheet.zip"
            )



        # =============================
        # FORMAT SWITCH (PDF vs DOCX)
        # =============================
        # =============================
        # FORMAT SWITCH (TXT / DOCX / PDF)
        # =============================
        return jsonify({"error": "Invalid worksheet type"}), 400

    except Exception as e:
        import traceback
        traceback.print_exc()
        if "429" in str(e):
            return jsonify({"error": "AI quota exceeded. Try again later."}), 429
        return jsonify({
        "error": str(e),
        "type": str(type(e))
    }), 500

def extract_images_from_pdf(file_storage):
    images = []
    pdf = fitz.open(stream=file_storage.read(), filetype="pdf")

    for page_index in range(len(pdf)):
        page = pdf[page_index]
        image_list = page.get_images(full=True)

        for img in image_list:
            xref = img[0]
            base_image = pdf.extract_image(xref)
            image_bytes = base_image["image"]

            img_pil = Image.open(io.BytesIO(image_bytes))
            images.append(img_pil)

    return images

def extract_images_from_docx(file_storage):
    images = []
    doc = Document(file_storage)

    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_bytes = rel.target_part.blob
            images.append(Image.open(io.BytesIO(image_bytes)))

    return images      

def requires_diagram(question_text):
    keywords = [
        "diagram", "figure", "graph", "number line",
        "draw", "construct", "triangle", "circle",
        "rectangle", "geometry"
    ]
    return any(k in question_text.lower() for k in keywords)

from docx.shared import Inches

def add_image_to_docx(doc, image):
    temp_path = "temp_img.png"
    image.save(temp_path)
    doc.add_picture(temp_path, width=Inches(3))

def create_image_with_diagram(content_lines, diagrams):
    img = Image.new("RGB", (1240, 1754), "white")
    draw = ImageDraw.Draw(img)
    
    from PIL import ImageFont

    try:
        font_body = ImageFont.truetype("DejaVuSans.ttf", 22)
    except:
        font_body = ImageFont.load_default()

    y = 40
    for i, line in enumerate(content_lines):
        draw.text((40, y), line, fill="black", font=font_body)
        y += 30

    y = 40
    for i, line in enumerate(content_lines):
        draw.text((40, y), line, fill="black", font=font_body)
        y += 30

        if i < len(diagrams):
            img.paste(diagrams[i], (40, y))
            y += diagrams[i].height + 20

    return img

def generate_triangle_diagram():
    img = Image.new("RGB", (400, 300), "white")
    draw = ImageDraw.Draw(img)

    points = [(200, 50), (80, 250), (320, 250)]
    draw.polygon(points, outline="black", width=3)

    return img


# --- Function to Create Database Tables AND SEED SAMPLE TESTS ---
def create_database(app_instance):
    with app_instance.app_context():
        db.create_all()
        print("Database tables ensured.")


# --- Main Execution Block ---
if __name__ == '__main__':
    # Check for required OAuth environment variables
    missing_vars = [var for var in ['GOOGLE_CLIENT_ID', 'GOOGLE_CLIENT_SECRET'] if not app.config.get(var)]
    if missing_vars:
        print("\n---!!! WARNING: Google OAuth Environment Variables Missing !!!---")
        print("Google login will not work until you set:")
        for var in missing_vars: print(f"  - {var}")
        print("-------------------------------------------------------\n")

    create_database(app) # Ensure database and tables are ready
    print("----------------------------------------------------")
    print("Flask server starting...")
    print(f"Database located at: {db_path}")
    print("Available Routes:")
    print("  - /                      (Main App - Public)")
    print("  - /login                 (Login Page)")
    print("  - /register              (Registration Page)")
    print("  - /logout                (Logout Action - Requires Login)")
    print("  - /login/google          (Initiate Google Login)")
    print("  - /auth/google/callback  (Google Callback)")
    print("  - /generate-worksheet    (API Endpoint - Requires Login)")
    print("  - /profile               (Profile Page - Requires Login)")
    print("  - /settings              (Settings Page - Requires Login)")
    print("  - /competitive-exams     (List Competitive Tests)")
    print("  - /job-exams             (List Job Screening Tests)")
    print("  - /mock-test/<test_id>   (Take a mock test)")
    print("  - /my-scores             (List your attempts)")
    print("  - /review/<attempt_id>   (Review an attempt)")
    print("----------------------------------------------------")
    app.run(debug=True, port=5000)
