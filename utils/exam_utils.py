import os
import re
import json
from datetime import datetime
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH



BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMP_DIR = os.path.join(BASE_DIR, "temp_files")
os.makedirs(TEMP_DIR, exist_ok=True)


def extract_json_from_ai(text):
    """
    Safely extract JSON array from Gemini output
    """
    if not text:
        return None

    # Try direct JSON
    try:
        return json.loads(text)
    except Exception:
        pass

    # Try extracting JSON inside text
    match = re.search(r"\[[\s\S]*\]", text)
    if match:
        try:
            return json.loads(match.group())
        except Exception:
            return None

    return None
  
def clean_ai_text(text):
    # Convert LaTeX fractions to a/b
    text = re.sub(r"\\frac\s*\{([^}]+)\}\{([^}]+)\}", r"\1/\2", text)

    if not text:
        return ""

    # ---------- LaTeX → Unicode Math ----------
    replacements = {
        # Greek
        "\\alpha": "α", "\\beta": "β", "\\gamma": "γ",
        "\\delta": "δ", "\\theta": "θ", "\\pi": "π",
        "\\lambda": "λ", "\\mu": "μ", "\\sigma": "σ",

        # Roots & powers
        "\\sqrt": "√",
        "^2": "²", "^3": "³",
        "^4": "⁴", "^5": "⁵",

        # Operators
        "\\times": "×",
        "\\cdot": "·",
        "\\pm": "±",
        "\\div": "÷",

        # Relations
        "\\le": "≤",
        "\\ge": "≥",
        "\\neq": "≠",
        "\\approx": "≈",

        # Brackets
        "\\left(": "(",
        "\\right)": ")",
        "\\left[": "[",
        "\\right]": "]",
        "\\left\\{": "{",
        "\\right\\}": "}",

        # Noise
        "$": "",
        "\\(": "",
        "\\)": "",
        "\\,": " ",
        "\\;": " ",
        "\\!": "",
        "\\": "",
    }

    for k, v in replacements.items():
        text = text.replace(k, v)

    # ---------- Fix remaining math formatting ----------
    # sqrt(x) → √x
    text = re.sub(r"√\(([^)]+)\)", r"√\1", text)

    # Remove double spaces
    text = re.sub(r"\s{2,}", " ", text)

    # Clean lines
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return "\n".join(lines)
  
def normalize_questions(questions_json):
    lines = []
    for i, q in enumerate(questions_json, 1):
        question = clean_ai_text(q["question"])
        lines.append(f"{i}) {question}")
        lines.append("")  # spacing between questions
    return "\n".join(lines)

def normalize_answers(text):
    if not text:
        return ""

    # Replace "1) answer 2) answer" → line breaks
    text = re.sub(r"\s*(\d+\))", r"\n\1", text)

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return "\n".join(lines)

# --- PDF Generation Class ---
class CustomPDF(FPDF):
    def __init__(self, title, sub_info, header_text="", footer_text=""):
        super().__init__()
        FONT_PATH = os.path.join(BASE_DIR, "fonts", "DejaVuSans.ttf")
        self.add_font("DejaVu", "", FONT_PATH, uni=True)
        self.worksheet_title = title
        self.sub_info = sub_info
        self.header_text = header_text
        self.footer_text = footer_text

    def header(self):
        self.set_font("DejaVu", "", 10)
        self.cell(0, 8, self.header_text, 0, 1, "C")

        self.set_font("DejaVu", "", 16)
        self.cell(0, 10, self.worksheet_title, 0, 1, "C")

        self.set_font("DejaVu", "", 10)
        sub_header_text = (
            f"Date: {self.sub_info['date']}   |   "
            f"Marks: {self.sub_info['marks']}   |   "
            f"Topic: {self.sub_info['sub-title']}"
        )
        self.cell(0, 8, sub_header_text, 0, 1, "C")
        self.ln(6)


    def footer(self):
        self.set_y(-15)
        self.set_font("DejaVu", "", 8)

        # --- Footer Left Text ---
        self.cell(0, 8, self.footer_text, 0, 0, "L")

        # --- Page Number (Right) ---
        self.cell(0, 8, f"Page {self.page_no()}", 0, 0, "R")

# --- File Creation Functions ---
def create_pdf(content, title, sub_title_info, filename="worksheet.pdf",
               header_text="", footer_text=""):

    uid = datetime.now().strftime("%Y%m%d%H%M%S")
    filename = f"Worksheet_{uid}.pdf"
    file_path = os.path.join(TEMP_DIR, filename)

    pdf = CustomPDF(
        title,
        sub_title_info,
        header_text=header_text,
        footer_text=footer_text
    )
    pdf.add_page()
    pdf.set_font("DejaVu", "", 12)

    pdf.multi_cell(0, 10, content)  # ✅ NO normalization here

    pdf.output(file_path)
    return file_path

def create_docx(content, title, sub_title_info, filename="worksheet.docx"):
    uid = datetime.now().strftime("%Y%m%d%H%M%S")
    filename = f"Worksheet_{uid}.docx"

    print("Generating DOCX...")

    doc = Document()

    # ---------- TITLE ----------
    title_p = doc.add_heading(title, level=1)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---------- SUB HEADER ----------
    sub = (
        f"Date: {sub_title_info['date']}   |   "
        f"Marks: {sub_title_info['marks']}   |   "
        f"Topic: {sub_title_info['sub-title']}"
    )
    sub_p = doc.add_paragraph(sub)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # ---------- CONTENT ----------
    for line in content.split("\n"):
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(8)

    # ---------- FOOTER ----------
    footer = doc.sections[0].footer
    footer_p = footer.paragraphs[0]
    footer_p.text = "MathGen | Generated for learning purposes"
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    file_path = os.path.join(TEMP_DIR, filename)
    doc.save(file_path)
    return file_path



def create_txt(content, title, sub_title_info, filename="worksheet.txt"):
    print("Generating TXT...")

    file_path = os.path.join(TEMP_DIR, filename)
    with open(file_path, "w", encoding="utf-8") as f:

        # ---------- HEADER ----------
        f.write("=" * 50 + "\n")
        f.write(f"{title}\n")
        f.write("=" * 50 + "\n\n")

        f.write(f"Date   : {sub_title_info.get('date', '')}\n")
        f.write(f"Time   : {sub_title_info.get('time', '')}\n")
        f.write(f"Marks  : {sub_title_info.get('marks', '')}\n")
        f.write(f"Topic  : {sub_title_info.get('sub-title', '')}\n")
        f.write("\n")

        # ---------- INSTRUCTIONS ----------
        f.write("-" * 50 + "\n")
        f.write("Instructions:\n")
        f.write("• Read each question carefully.\n")
        f.write("• Show your working clearly.\n")
        f.write("• Write answers in the space provided.\n")
        f.write("-" * 50 + "\n\n")

        # ---------- QUESTIONS ----------
        f.write(content.strip())
        f.write("\n\n")

        # ---------- FOOTER ----------
        f.write("-" * 50 + "\n")
        f.write("End of Worksheet\n")
        f.write("-" * 50 + "\n")

    return file_path
